"""``DocxFormat`` — Microsoft Word ``.docx`` (Office Open XML) encoder/decoder.

Reads and writes use ``python-docx``. ``.docx`` is a zipped XML bundle:
random access requires the full archive, so this is a
:class:`BatchFileFormat`.

Each paragraph in the document maps to one record with ``index``,
``text`` and ``style`` fields. Encoding is the inverse: each record adds
one paragraph with the given style (or the default style when omitted).

Security: pirn does not sandbox ``python-docx``. Malformed archives may
trigger upstream library bugs. Treat untrusted payloads accordingly.

Install: ``pip install pirn[docx]``.
"""

from __future__ import annotations

import io
from collections.abc import Iterable, Mapping
from typing import Any

from pirn.domains.connectors.file_formats.batch_file_format import (
    BatchFileFormat,
)


class DocxFormat(BatchFileFormat):
    """Whole-file DOCX encoder/decoder."""

    def __init__(self, paragraph_separator: str = "\n") -> None:
        if not isinstance(paragraph_separator, str):
            raise TypeError(
                "DocxFormat: paragraph_separator must be a string, "
                f"got {type(paragraph_separator).__name__}"
            )
        self._paragraph_separator = paragraph_separator

    @property
    def name(self) -> str:
        return "docx"

    @property
    def paragraph_separator(self) -> str:
        return self._paragraph_separator

    async def _decode_full(self, payload: bytes) -> Iterable[Mapping[str, Any]]:
        docx = self._load_docx()
        document = docx.Document(io.BytesIO(payload))
        records: list[Mapping[str, Any]] = []
        for index, paragraph in enumerate(document.paragraphs):
            style_name = ""
            if paragraph.style is not None:
                style_name = paragraph.style.name or ""
            records.append(
                {
                    "index": index,
                    "text": paragraph.text,
                    "style": style_name,
                }
            )
        return records

    async def _encode_full(self, records: Iterable[Mapping[str, Any]]) -> bytes:
        docx = self._load_docx()
        document = docx.Document()
        for record in records:
            if "text" not in record:
                raise ValueError(
                    "DocxFormat: record missing required 'text' field "
                    f"— got keys {list(record.keys())}"
                )
            text = record["text"]
            if not isinstance(text, str):
                raise TypeError(f"DocxFormat: 'text' must be a string, got {type(text).__name__}")
            style = record.get("style")
            if style is not None and not isinstance(style, str):
                raise TypeError(
                    "DocxFormat: 'style' must be a string when "
                    f"provided, got {type(style).__name__}"
                )
            if style:
                document.add_paragraph(text, style=style)
            else:
                document.add_paragraph(text)
        buf = io.BytesIO()
        document.save(buf)
        return buf.getvalue()

    @staticmethod
    def _load_docx() -> Any:
        try:
            import docx
        except ImportError as exc:
            raise ImportError(
                "DocxFormat requires python-docx. Install with `pip install pirn[docx]`."
            ) from exc
        return docx
